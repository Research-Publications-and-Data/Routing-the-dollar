"""
Build the Online Supplement for Routing the Dollar v42.

Assembles content from paper_v25.md, data/processed/, and media/ into
a formatted .docx file that resolves all 11 supplement references in
the main paper.
"""
import os, sys, re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
MEDIA = ROOT / "media"
DATA_PROC = ROOT / "data" / "processed"
OUTPUT = ROOT / "Routing_the_Dollar_Supplement_v42.docx"

# ── Styling helpers ──────────────────────────────────────────

def set_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    return h

def add_para(doc, text, bold=False, italic=False, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p

def add_source(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p

def add_exhibit(doc, filename, caption, width=6.0):
    path = MEDIA / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_source(doc, caption)
    else:
        add_para(doc, f"[Exhibit image not found: {filename}]", italic=True)
        add_source(doc, caption)

def add_simple_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def add_panel_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)

# ── Language sync ────────────────────────────────────────────

STALE_REPLACEMENTS = [
    ('fully endogenous', 'shared exposure to money-market conditions rather than a policy reaction function'),
    ('Rate cuts predict', 'Declines in Treasury bill yields predict'),
    ('consists largely of peer-to-peer', 'appears predominantly P2P-like based on the heuristics and bounds described below'),
    ('institutionalizes this dynamic', 'reinforces this issuer-side channel under the current legal design'),
    ('Annual on-chain transfer volume', 'On-chain stablecoin transfer volume across nine blockchains totals'),
    ('de facto compliance screening', 'subjecting the routed subset of flows to their compliance programs'),
    ('across borders, protocols', 'across blockchains, protocols'),
    ('the determinant was entity', 'salient variation'),
    ('cross-gateway bridges halved', 'cross-gateway counterparties roughly halved'),
]

def language_sync(text):
    count = 0
    for stale, correct in STALE_REPLACEMENTS:
        n = text.count(stale)
        if n > 0:
            text = text.replace(stale, correct)
            count += n
            print(f"  REPLACED: '{stale}' ({n}x)")
    # stress event -> stress episode (but not "stress episode" already)
    n = 0
    # Use regex to avoid replacing inside "stress episode"
    import re
    matches = list(re.finditer(r'stress event', text))
    for m in reversed(matches):
        text = text[:m.start()] + 'stress episode' + text[m.end():]
        n += 1
    if n:
        print(f"  REPLACED: 'stress event' -> 'stress episode' ({n}x)")
        count += n
    # confirms this pattern -> is consistent with this pattern
    # But NOT "confirms this feedback interpretation" or confirms in statistical contexts
    for match in list(re.finditer(r'confirms this pattern', text)):
        text = text[:match.start()] + 'is consistent with this pattern' + text[match.end():]
        count += 1
        print(f"  REPLACED: 'confirms this pattern' (1x)")
    return text, count

def style_check(text):
    em_dashes = text.count('\u2014')
    authors_plural = len(re.findall(r"Authors'|Authors\u2019", text))
    author_singular = len(re.findall(r"Author's|Author\u2019s", text))
    stress_event = text.count('stress event')
    return {
        'em_dashes': em_dashes,
        'authors_plural': authors_plural,
        'author_singular': author_singular,
        'stress_event': stress_event,
    }

def fix_style(text):
    # Remove em-dashes by rephrasing (use comma/semicolon, not dash substitution)
    text = text.replace('\u2014', ', ')
    # Fix Authors' -> Author's
    text = text.replace("Authors'", "Author's")
    text = text.replace("Authors\u2019", "Author\u2019s")
    return text

# ── Section builders ─────────────────────────────────────────

def build_s1_intro(doc):
    add_heading(doc, "Introduction", level=1)
    add_para(doc, (
        "This is the companion document posted with \"Routing the Dollar: Gateway Infrastructure, "
        "Monetary Policy Transmission, and the Dollar's International Functions in Digital Markets,\" "
        "containing tables, exhibits, and extended analyses omitted "
        "from the main text for length. References of the form \"(online supplement)\" in the main "
        "text correspond to sections below."
    ))
    add_para(doc, (
        "Appendix lettering follows the working-paper series. The main paper contains Appendices "
        "B, C, D, E, F, G, H, and K. This supplement provides the content for omitted appendix "
        "letters (A, J) and additional appendices (on tokenized fund assessment and the credit "
        "migration channel) removed from the main text for length, along with the detailed econometric "
        "tables, robustness analyses, and extended exhibits referenced throughout the paper. "
        "Appendix I (Legislative Framework and the Control Layer) appears as Appendix K in the "
        "main paper."
    ))

def build_s2_exhibit_inventory(doc):
    add_heading(doc, "Appendix A: Complete Exhibit Inventory", level=1)
    add_para(doc, (
        "Table A1 inventories all exhibits in the main paper and supplement, with key findings."
    ))
    headers = ["#", "Title", "Key Finding"]
    rows = [
        ["1", "VECM IRFs (bootstrap 90% CI)", "Fed assets shock -> supply: significant at weeks 8, 12, 26"],
        ["2", "FEVD: Stablecoin Supply", "Fed + ON RRP: 13.9% of supply variance at 26 weeks"],
        ["3", "Persistence Profile", "Half-life 5 weeks; 6-8 week oscillation (FOMC cycle)"],
        ["4", "Yield Spread vs. Supply Growth", "Granger F = 11.30, p = 0.001"],
        ["5", "Correlation Matrix", "Fed assets r = -0.94; fed funds r = -0.89; SOFR r = -0.87"],
        ["6", "Rolling Correlations (90/180-day)", "Sign reversal: +0.93 (QE) -> ~0 (tightening) -> -0.98 (easing)"],
        ["7", "Supply vs. Fed Funds Rate", "Negative co-movement across full sample"],
        ["8", "Supply vs. ON RRP", "$2.4T drawdown coincides with supply expansion"],
        ["9", "FOMC Event Study", "Directional but not significant after Bonferroni"],
        ["10", "Use-Case Decomposition", "P2P 45.7%; DeFi 33.2%; CEX 8.3%; payments 0.1%"],
        ["11", "Market Capitalization Panels", "USDT dominance rising; BUSD eliminated"],
        ["12", "Gateway Volume by Tier", "T2 55% gross/50% net; T1 41%/47%; duopoly"],
        ["13", "SVB Crisis Dynamics", "Entity-specific reallocation; Tier 3 absorbed 66-69%"],
        ["14", "SVB Flow Retention", "CLII-retention gradient: r = -0.34 (SVB), -0.47 (BUSD)"],
        ["15", "Gateway Concentration (HHI)", "Mean 5,021; range 3,800-7,200"],
        ["16", "Counterparty Network", "Volume doubled; unique counterparties -25%"],
        ["17", "Network Topology Time Series", "Cross-gateway bridges: 5.8% -> 2.6%"],
        ["18", "Multi-Chain Control Layer", "Tron: 0% T1; Ethereum: 41% T1; Solana: 11% T1"],
        ["19", "CLII-Retention Gradient", "Continuous: Q1 6.27x, Q4 0.61x retention"],
        ["20", "Trivariate Robustness", "Survives SOFR addition; 10-year yield fails"],
        ["21", "Jackknife Stability", "No single entity drives aggregate results"],
        ["22", "CLII Weight Sensitivity", "18/19 entities invariant across 5 weighting schemes"],
        ["23", "Event-Time Placebo", "SVB: 94th pctile of 50 placebos; nadir z = -2.7sigma"],
        ["24", "Coverage Sensitivity", "T1 share: 43.5% -> 26.3% at 25% reclassification"],
        ["25", "DEX Volume vs. Supply", "Co-trending but DEX higher volatility"],
        ["26", "IRF Lag Robustness", "Lag 8 vs. lag 4 qualitatively similar"],
        ["27", "ON RRP Decomposition", "$2.4T -> $9.6B drawdown visualization"],
        ["C2", "Placebo T1 Swings (orig.)", "SVB 23.3 pp > all 50 placebos"],
        ["C2b", "Event-Time Stress Placebo", "SVB 2.7sigma below mean; below 5th pctile 2/15 days"],
        ["C8", "No-Freeze CLII Robustness", "0/19 tier changes; max shift 0.11"],
        ["C9", "Day-of-Week Structure", "Structural weekend dip: 13.9 pp"],
        ["C10", "Placebo (Expanded Registry)", "SVB swing 51.8 pp; 94th percentile"],
        ["SB-1", "Unlabeled Address Degree", "Labeled median degree 261 vs. unlabeled 41"],
        ["SB-2", "Transfer Size Distribution", "Labeled $1M+ share: 84.3% vs. 66.7%"],
    ]
    add_simple_table(doc, headers, rows)
    add_source(doc, "Table A1. Complete exhibit inventory with key findings.")

def build_s3_tables_f1_f3(doc):
    add_heading(doc, "Tables F1--F3: Consolidated Econometric Results", level=1)

    # Table F1 (= v25 Table E1)
    add_heading(doc, "Table F1. Consolidated Econometric Results", level=2)

    add_panel_header(doc, "Panel A: Cointegration")
    headers = ["Test", "Statistic", "Result"]
    rows = [
        ["Johansen trace (rank <= 0)", "30.68 [cv95 = 29.80]*", "Reject: rank = 1"],
        ["Johansen lambda-max (rank <= 0)", "23.71 [cv95 = 21.13]*", "Reject"],
        ["Cointegrating rank", "1", "--"],
        ["Engle-Granger (bivariate)", "-2.92 [p = 0.13]", "Fail to reject"],
        ["System", "Log Fed assets, log ON RRP, log stablecoin supply", "--"],
        ["Sample", "Weekly, Feb 2023 to Jan 2026, T = 155, lag = 8 (AIC)", "--"],
    ]
    add_simple_table(doc, headers, rows)

    add_panel_header(doc, "Panel B: VECM Adjustment Speeds (alpha)")
    headers = ["Variable", "alpha", "SE", "t-stat", "p-value"]
    rows = [
        ["Fed assets (WSHOMCB)", "0.003", "0.0004", "6.256", "<0.001***"],
        ["ON RRP (RRPONTSYD)", "-0.157", "0.135", "-1.166", "0.244"],
        ["Stablecoin supply", "-0.004", "0.002", "-1.589", "0.112"],
    ]
    add_simple_table(doc, headers, rows)

    add_para(doc, "Weak exogeneity LR tests (chi-squared(1)):", bold=True, size=10)
    headers = ["Variable", "LR statistic", "p-value", "Decision"]
    rows = [
        ["Fed assets", "15.28", "<0.001", "Reject"],
        ["ON RRP", "14.00", "<0.001", "Reject"],
        ["Stablecoin supply", "5.38", "0.020", "Reject at 5%"],
    ]
    add_simple_table(doc, headers, rows)

    add_panel_header(doc, "Panel C: FOMC Event Study (Raw / Abnormal)")
    headers = ["Classification", "t+5 raw", "t+5 abnormal", "t+10 raw", "t+10 abnormal"]
    rows = [
        ["Dovish (n = 5)", "+0.63% [p = 0.20]", "+0.13% [p = 0.79]", "+1.66% [p = 0.09]", "+0.65% [p = 0.50]"],
        ["Neutral (n = 9)", "+0.33% [p = 0.008]**", "-0.29% [p = 0.098]", "+0.93% [p = 0.06]", "-0.31% [p = 0.36]"],
        ["Hawkish (n = 5)", "-0.07% [p = 0.57]", "+0.19% [p = 0.36]", "-0.48% [p = 0.07]", "+0.04% [p = 0.85]"],
    ]
    add_simple_table(doc, headers, rows)

    add_panel_header(doc, "Panel D: Deposit Displacement")
    headers = ["Measure", "Statistic"]
    rows = [["Differenced r (weekly)", "0.12 [p = 0.024]*"]]
    add_simple_table(doc, headers, rows)

    add_source(doc, (
        "Table F1. Consolidated econometric results. Panel B: only Fed assets alpha is individually "
        "significant; ON RRP and stablecoin supply carry error-correcting signs but are imprecisely "
        "estimated. Panel C: abnormal returns subtract the trailing 10-day average daily supply growth "
        "x horizon; no raw or abnormal result survives Bonferroni correction. "
        "Source: Author's calculations using FRED and DefiLlama data."
    ))

    # Table F2 (= v25 Table E2)
    doc.add_page_break()
    add_heading(doc, "Table F2. Comprehensive Robustness and Identification Summary", level=2)

    add_panel_header(doc, "Panel A: Cointegration and Identification")
    headers = ["Test", "Statistic", "Result"]
    rows = [
        ["Johansen trace (lag 8)", "30.68 (cv95 = 29.80)", "Reject: rank = 1"],
        ["Lag sensitivity", "3/5 lags pass (8, 10, 12)", "Robust"],
        ["KPSS (all variables)", "I(1) confirmed", "Consistent"],
        ["ADF-KPSS supply (differenced)", "ADF p = 0.002; KPSS p = 0.040", "Mild disagreement"],
        ["Placebo: BTC", "0/5 lags cointegrate", "Clean falsification"],
        ["Placebo: ETH", "0/5 lags cointegrate", "Clean falsification"],
        ["Cross-stablecoin: USDT", "2/5 lags pass (10, 12)", "Supportive"],
        ["Cross-stablecoin: USDC", "2/5 lags pass (4, 6)", "Supportive"],
        ["USDT alpha (error correction)", "-0.005 (p = 0.005)", "Individually significant"],
        ["Beta stability (Sep 2024 break)", "Bootstrap p > 0.05 all elements", "Stable"],
        ["Gregory-Hansen (C/S model)", "ADF* = -4.28 (cv 5% = -5.96)", "Not significant"],
        ["Yield spread Granger (lag 1)", "F = 11.30 (p = 0.001)", "Significant"],
        ["Supply -> spread Granger (lag 2)", "F = 9.40 (p < 0.001)", "Bidirectional"],
        ["ON RRP Granger (lag 1)", "F = 4.83 (p = 0.028)", "Significant"],
        ["ON RRP Granger (lag 2)", "F = 3.91 (p = 0.022)", "Significant"],
        ["Supply -> ON RRP (lag 5)", "F = 4.13 (p = 0.002)", "Bidirectional"],
        ["FEVD at 26 weeks", "Fed + RRP: 13.9% of supply variance", "Moderate policy share"],
        ["Persistence profile (Pesaran-Shin)", "Smoothed half-life: 5 weeks", "Rapid equilibrium restoration"],
        ["FOMC volume ratio", "0.69x (t = -3.83, p < 0.001)", "Significant (parametric)"],
        ["FOMC volume Mann-Whitney", "U = 9,542 (p = 0.614)", "Not significant"],
        ["Flow retention vs CLII (SVB)", "r = -0.34 (p = 0.27; n = 12)", "Directionally consistent"],
        ["Flow retention vs CLII (BUSD)", "r = -0.47 (p = 0.13; n = 12)", "Directionally consistent"],
        ["Flow retention pooled (SVB + BUSD)", "r = -0.30 (p = 0.16; n = 24)", "Direction holds"],
    ]
    add_simple_table(doc, headers, rows)

    add_panel_header(doc, "Panel E: Gateway Tier Correlations (Expanded Registry, Weekly, n = 158)")
    headers = ["Relationship", "Statistic", "Interpretation"]
    rows = [
        ["Tier 1 vs Fed Assets (levels)", "r = -0.56 (p < 0.001)", "Long-run co-trend"],
        ["Tier 2 vs Fed Assets (levels)", "r = -0.45 (p < 0.001)", "Moderate co-movement"],
        ["Tier 3 vs Fed Assets (levels)", "r = +0.37 (p < 0.001)", "Opposite direction (DeFi growth)"],
        ["Tier 1 vs Fed Assets (first diff.)", "r = +0.05 (p = 0.49)", "Not significant"],
        ["Tier 2 vs Fed Assets (first diff.)", "r = +0.12 (p = 0.14)", "Not significant"],
        ["Tier 3 vs Fed Assets (first diff.)", "r = +0.07 (p = 0.40)", "Not significant"],
    ]
    add_simple_table(doc, headers, rows)

    add_panel_header(doc, "Panel F: SVB Counterparty Network (Nansen, 15 Gateways, 3 Windows)")
    headers = ["Metric", "Pre / Stress / Post", "Interpretation"]
    rows = [
        ["Network density", "0.085 / 0.085 / 0.086", "Topology stable; capacity fragmented"],
        ["Market-maker volume change (stress vs pre)", "-29% aggregate", "Cumberland -79%, Wintermute +118%"],
        ["Wintermute gateway connections", "6 / 6 / 6", "Maintained all; added Circle"],
        ["Cumberland gateway connections", "6 / 3 / 6", "Lost Curve, Coinbase, Paxos during stress"],
        ["MEV bot share of Curve 3pool", "24% pre -> 51% stress", "Replaced institutional intermediaries"],
        ["Core triangle (Binance-Circle-Tether)", "44% / 42% / 57%", "Post-crisis consolidation"],
    ]
    add_simple_table(doc, headers, rows)

    add_panel_header(doc, "Panel G: BUSD Wind-Down Counterparty Network (Nansen, 15 Gateways, 3 Windows)")
    headers = ["Metric", "Pre / Early / Late", "Interpretation"]
    rows = [
        ["Wintermute volume", "$4.7B / $12.0B / $24.4B", "5x surge; mediated BUSD->USDT"],
        ["Paxos as Binance counterparty", "$6.8B / $4.3B / $0.8B", "88% decline"],
        ["Tether Treasury counterparty count", "15 / 64 / 500", "Broadening USDT mint"],
        ["Total network volume", "$410B / $507B / $457B", "Rose then stabilized"],
    ]
    add_simple_table(doc, headers, rows)

    add_panel_header(doc, "Panel H: Layer 2 Gateway Tier Composition (Nansen)")
    headers = ["Chain", "T1", "T2", "T3", "Note"]
    rows = [
        ["Base", "0%", "0%", "100%", "Morpho Blue dominates"],
        ["Arbitrum", "0%", "46%", "54%", "Binance + Hyperliquid/Aave/GMX"],
        ["Ethereum (comparison)", "44%", "53%", "3%", "L2s do not inherit parent tier structure"],
    ]
    add_simple_table(doc, headers, rows)

    add_panel_header(doc, "Panel I: Monthly Market-Maker Concentration (Nansen, 15 Gateways, 35 Months)")
    headers = ["Metric", "Value", "Interpretation"]
    rows = [
        ["Wintermute monthly volume", "$2.2B (Feb 2023) -> $51.8B (Aug 2025)", "24x growth; 1.4% -> 19.9%"],
        ["Cumberland monthly volume", "$5.2B (Feb 2023) -> $0.2B (Dec 2025)", "96% decline"],
        ["Wintermute/Cumberland ratio", "0.4x -> 42x", "Market-maker layer concentrating"],
        ["Cross-gateway counterparties", "5.8% (Feb 2023) -> 2.6% (Dec 2025)", "Shared counterparties declining"],
    ]
    add_simple_table(doc, headers, rows)

    add_panel_header(doc, "Panel J: Solana Gateway Tier Composition (Dune Analytics, 14 Entities)")
    headers = ["Metric", "Value", "Interpretation"]
    rows = [
        ["Tier 1 volume share", "11.1% ($49.7B of $446.7B)", "Coinbase $25.0B + Circle $24.8B"],
        ["Tier 2 volume share", "88.9% ($396.9B)", "Binance $278B dominates"],
        ["Tier 3 DEX volume", "$4.2T synthetic", "Jupiter/Raydium/Orca; not comparable"],
        ["Solana vs Ethereum T1 share", "11% vs 41%", "Exchange-dominated, DeFi-first"],
    ]
    add_simple_table(doc, headers, rows)

    add_source(doc, (
        "Table F2. Comprehensive robustness and identification summary. See Section IV.A.1 of the "
        "main paper for detailed discussion. Source: Author's calculations."
    ))

    # Table F3 (= v25 Table E3)
    add_heading(doc, "Table F3. Gateway Volume on FOMC Days", level=2)
    headers = ["Classification", "Volume Ratio", "Interpretation"]
    rows = [
        ["Hawkish", "0.50x", "Lowest activity; uncertainty peak"],
        ["Dovish", "0.72x", "Moderate suppression"],
        ["Neutral", "0.77x", "Closest to baseline"],
    ]
    add_simple_table(doc, headers, rows)
    add_source(doc, (
        "Table F3. Gateway volume on FOMC days relative to non-FOMC baseline. Daily stablecoin "
        "gateway transfer volume is 31 percent lower on FOMC announcement days (Welch's t = -3.83, "
        "p < 0.001), though the non-parametric Mann-Whitney test does not confirm (U = 9,542, p = 0.614). "
        "Source: Author's calculations using Dune Analytics data."
    ))

def build_s4_fomc_extended(doc):
    add_heading(doc, "Extended FOMC Event Study Analysis", level=1)
    add_para(doc, (
        "Classification-level detail. In the dovish case, the raw t+10 mean (+1.66%) is pulled "
        "substantially above the median (+1.14%) by the November 2024 rate cut, which coincided "
        "with a broader risk-on episode in crypto markets; the abnormal t+10 mean (+0.65%) is not "
        "significant (p = 0.50). The hawkish category shows the most notable abnormal result at the "
        "shortest horizon: hawkish announcements produce an abnormal t+1 supply decline of -0.039 "
        "percent (p = 0.027), suggesting that the immediate market reaction to tightening signals is "
        "the most cleanly identified policy effect in our sample, though the small sample (n = 5) "
        "warrants caution. With 12 hypothesis tests across three classifications and four horizons, "
        "neither the raw neutral t+5 result (p = 0.008) nor the hawkish abnormal t+1 result (p = 0.027) "
        "survives strict Bonferroni correction (adjusted alpha = 0.004); both survive at the 10 percent "
        "family-wise level."
    ))
    add_para(doc, (
        "FOMC-day volume effects. Daily stablecoin gateway transfer volume is 31 percent lower on "
        "FOMC announcement days relative to non-FOMC days (Welch's t = -3.83, p < 0.001), though "
        "the non-parametric Mann-Whitney test does not confirm this finding (U = 9,542, p = 0.614). "
        "The parametric/non-parametric disagreement warrants caution, but the directional result "
        "suggests stablecoin markets have become integrated enough with the traditional financial "
        "system that participants reduce activity around scheduled policy announcements."
    ))

def build_s5_unit_root(doc):
    add_heading(doc, "Unit Root Test Results (ADF and KPSS)", level=1)
    add_para(doc, (
        "Table F2, Panel A reports the joint ADF and KPSS results confirming that all three VECM "
        "variables (log Fed assets, log ON RRP, log stablecoin supply) are I(1). The ADF test rejects "
        "the unit root null in first differences for all series, while KPSS fails to reject stationarity "
        "in first differences, jointly confirming that each series is integrated of order one."
    ))
    add_para(doc, (
        "The ADF test on differenced Fed total assets is borderline (p = 0.19), but KPSS confirms "
        "stationarity (0.386, p = 0.083). The mild disagreement for stablecoin supply (ADF p = 0.002 "
        "vs. KPSS p = 0.040) does not affect the joint I(1) classification. The Johansen cointegration "
        "test is valid for I(1) series regardless of borderline individual test results."
    ))

def build_s6_weak_exogeneity(doc):
    add_heading(doc, "Johansen Weak Exogeneity Test", level=1)
    add_para(doc, (
        "Table F1, Panel B reports the weak exogeneity LR test results. The Johansen weak exogeneity "
        "test rejects alpha = 0 for all three variables in the VECM system: Fed assets (LR = 15.28, "
        "p < 0.001), ON RRP (LR = 14.00, p < 0.001), and stablecoin supply (LR = 5.38, p = 0.020). "
        "All three variables participate in error correction, reflecting shared exposure to money-market "
        "conditions rather than a policy reaction function."
    ))
    add_para(doc, (
        "This does not imply the Fed responds to stablecoin supply, an economically implausible reading. "
        "Rather, the Fed's administrative QT pace creates its own mean-reverting dynamics (alpha = +0.003, "
        "positive), while ON RRP and stablecoin supply carry error-correcting signs (alpha < 0), absorbing "
        "the adjustment. The apparent contradiction between the t-test insignificance of ON RRP's alpha "
        "(t = -1.166, p = 0.244) and the LR test rejection (LR = 14.00, p < 0.001) arises because the "
        "LR test evaluates the full system restriction (setting alpha to zero and re-estimating the entire "
        "cointegrating vector), while the t-test conditions on a fixed beta. The eigenvalue shifts by "
        "57.1 percent under the alpha = 0 restriction, confirming ON RRP's structural participation in "
        "equilibrium correction."
    ))

def build_s7_fevd_persistence(doc):
    add_heading(doc, "FEVD and Persistence Profile Detail", level=1)

    add_heading(doc, "Forecast Error Variance Decomposition", level=2)
    add_para(doc, (
        "The FEVD shows a slowly rising share of stablecoin supply forecast error variance attributable "
        "to Fed policy shocks, consistent with a gradual transmission mechanism rather than instantaneous "
        "pass-through."
    ))
    headers = ["Horizon (weeks)", "Fed Assets Share", "ON RRP Share", "Supply Own Share"]
    rows = [
        ["1", "0.8%", "2.8%", "96.4%"],
        ["4", "1.5%", "3.8%", "94.7%"],
        ["12", "5.2%", "5.1%", "89.7%"],
        ["26", "8.8%", "5.1%", "86.1%"],
    ]
    add_simple_table(doc, headers, rows)
    add_para(doc, (
        "Combined Fed + ON RRP share rises monotonically from 3.6% at 1 week to 13.9% at 26 weeks, "
        "indicating that policy shocks explain a modest but growing share of supply variance over "
        "the medium term."
    ))

    add_heading(doc, "Persistence Profile", level=2)
    add_para(doc, (
        "The persistence profile (Pesaran and Shin 1996) measures the speed at which the cointegrating "
        "equilibrium is restored following a system-wide shock. The smoothed half-life is approximately "
        "5 weeks, with 58 percent absorbed within one quarter. The profile exhibits oscillations at a "
        "6-8 week cycle consistent with the FOMC meeting schedule, explaining why FOMC announcement "
        "effects are individually weak while the long-run equilibrium is strong: transmission operates "
        "over weeks to months rather than days."
    ))
    add_source(doc, "Source: Author's calculations using FRED and DefiLlama data.")

def build_s8_svb_all_entities(doc):
    add_heading(doc, "SVB Gateway Flow Retention: All 12 Entities", level=1)
    add_para(doc, (
        "Table 6 in the main paper shows the top 10 gateways by volume. The complete 12-entity table "
        "is below. Flow retention is defined as the ratio of stress-period volume share to pre-stress "
        "baseline volume share; values above 1.0 indicate the gateway attracted disproportionate flows "
        "during SVB stress."
    ))
    headers = ["Gateway", "Tier", "CLII", "Baseline Share", "Stress Share", "Retention Ratio"]
    rows = [
        ["Circle Treasury", "1", "0.92", "34.2%", "23.8%", "1.39x"],
        ["Coinbase", "1", "0.85", "8.1%", "10.9%", "2.69x"],
        ["Paxos", "1", "0.88", "4.3%", "5.5%", "2.57x"],
        ["Gemini", "1", "0.82", "2.1%", "2.3%", "2.21x"],
        ["BitGo", "1", "0.80", "0.8%", "2.7%", "6.80x"],
        ["Tether Treasury", "2", "0.45", "15.2%", "13.5%", "1.77x"],
        ["Binance", "2", "0.38", "18.4%", "16.1%", "1.75x"],
        ["Kraken", "2", "0.58", "3.6%", "3.9%", "2.14x"],
        ["OKX", "2", "0.40", "2.8%", "2.5%", "1.79x"],
        ["Uniswap V3", "3", "0.15", "4.1%", "11.1%", "5.42x"],
        ["Curve 3pool", "3", "0.18", "1.2%", "7.4%", "12.35x"],
        ["Aave V3", "3", "0.20", "0.5%", "1.9%", "7.74x"],
    ]
    add_simple_table(doc, headers, rows)
    add_para(doc, (
        "The CLII-retention gradient is negative for both SVB (r = -0.34, p = 0.27, n = 12) and BUSD "
        "(r = -0.47, p = 0.13, n = 12) stress episodes, though neither reaches statistical significance "
        "at the 12-entity sample size. The pooled estimate (r = -0.30, p = 0.16, n = 24) confirms the "
        "directional pattern holds across both episodes."
    ))
    add_source(doc, (
        "Source: Author's calculations using Dune Analytics gateway transfer data. "
        "Baseline: 30-day pre-stress average. Stress: March 10-14, 2023 (SVB)."
    ))

def build_s9_placebo(doc):
    add_heading(doc, "Placebo Analysis: Full Results", level=1)
    add_para(doc, (
        "To test whether the SVB-week Tier 1 share movements are statistically exceptional, we "
        "compute event-time-aligned Tier 1 share trajectories for 50 randomly selected non-event "
        "windows over a [-7, +7] day horizon centered on each placebo date."
    ))

    add_heading(doc, "Original 12-Address Registry (Exhibit C2)", level=2)
    add_para(doc, (
        "Using the original 12-address registry, the SVB-week swing (23.3 pp) exceeds all 50 "
        "randomly sampled non-event windows (max placebo: 21.9 pp; mean: 13.4 pp)."
    ))
    add_exhibit(doc, "exhibit_placebo_t1_share.png",
        "Exhibit C2. Placebo analysis of Tier 1 share swings (original 12-address registry). "
        "Source: Author's calculations using Dune Analytics gateway share data.")

    add_heading(doc, "Expanded 51-Address Registry (Exhibit C10)", level=2)
    add_para(doc, (
        "Under the expanded 51-address registry, the SVB trajectory's 15-day swing of 51.8 "
        "percentage points ranks at the 94th percentile of placebo swings, its nadir falls 2.7 "
        "standard deviations below the placebo mean, and the trajectory remains below the 5th "
        "percentile band for 2 of 15 event-time days."
    ))
    add_exhibit(doc, "exhibit_placebo_expanded.png",
        "Exhibit C10. Placebo analysis with expanded 51-address registry. Gray: 50 non-event "
        "windows; blue band: 5th-95th percentile; dashed: placebo mean; red: SVB trajectory. "
        "Source: Author's calculations using Dune Analytics data.")

    add_source(doc, (
        "The multi-day swing exceeds all but 3 of 50 placebo windows, confirming the SVB "
        "stress episode produced statistically exceptional gateway reallocation."
    ))

def build_s10_day_of_week(doc):
    add_heading(doc, "Day-of-Week Tier 1 Structure", level=1)
    add_para(doc, (
        "The expanded 51-address registry reveals a structural weekend dip in Tier 1 share "
        "(weekday mean 44.8 percent, weekend mean 30.9 percent, gap 13.9 pp), driven primarily "
        "by Binance's 24/7 operations increasing Tier 2's weekend share. This structural pattern "
        "means that the SVB weekend nadir of 13.2 percent (March 12) cannot be evaluated against "
        "weekday baselines; the relevant comparison is against normal weekend variation."
    ))
    add_para(doc, (
        "At the weekend level, the SVB dip falls 0.85 standard deviations below the normal "
        "weekend mean, not statistically exceptional. Day-over-day changes provide additional "
        "signal: the SVB Friday-to-Saturday drop (-32.7 pp) exceeds the normal Friday-to-Saturday "
        "mean (-14.9 pp) by 1.4 sigma, and the Saturday-to-Sunday continuation (-16.6 pp) exceeds "
        "its normal mean (-0.2 pp) by 1.3 sigma. Neither is individually significant; the combined "
        "two-day trajectory from 62.5 to 13.2 percent is what the placebo analysis identifies as "
        "exceptional."
    ))
    add_exhibit(doc, "exhibit_weekend_dayofweek.png",
        "Exhibit C9. Day-of-week structure and SVB weekend effects (expanded 51-address registry). "
        "Source: Author's calculations using Dune Analytics data.")

def build_s11_no_freeze(doc):
    add_heading(doc, "No-Freeze CLII Robustness", level=1)
    add_para(doc, (
        "The freeze/blacklist capability dimension (20 percent weight) partially proxies for "
        "centralization rather than regulatory compliance per se: centralized entities can freeze "
        "assets while permissionless protocols cannot, regardless of regulatory intent. To test "
        "whether this conflation drives the tier classification, we recompute CLII scores dropping "
        "the freeze dimension entirely and redistributing its weight proportionally across the "
        "remaining four dimensions (License 31.25 percent, Transparency 25 percent, Compliance "
        "25 percent, Sanctions 18.75 percent)."
    ))
    add_para(doc, (
        "Zero of 19 entities change tier classification. The maximum absolute CLII shift is 0.11 "
        "(Tether, whose high freeze score of 0.90 contrasts with low scores on other dimensions); "
        "Tier 1 entities shift by less than 0.02. The tier structure reflects licensing, transparency, "
        "compliance infrastructure, and sanctions screening, not freeze capability alone."
    ))
    add_exhibit(doc, "exhibit_clii_nofreeze_robustness.png",
        "Exhibit C8. CLII no-freeze robustness. Dumbbell chart showing baseline CLII score (left dot) "
        "and no-freeze CLII score (right dot) for all 19 entities, with tier boundaries (0.75, 0.30) "
        "as vertical lines. Source: Author's calculations.")

def build_s12_unlabeled(doc):
    add_heading(doc, "Unlabeled Address Characterization", level=1)
    add_para(doc, (
        "Our labeled set is a sample of gateway addresses, not the universe. The principal threat "
        "is that unlabeled volume (72.3 percent of Ethereum USDC/USDT transfers) is systematically "
        "gateway-like rather than P2P-like, which would bias tier shares downward for whichever tier "
        "the unlabeled gateways belong to."
    ))
    add_para(doc, (
        "Transfer-level behavioral analysis comparing the 51 labeled gateway addresses against the "
        "top 500 unlabeled addresses by volume (6-month sample, Jul 2024-Jan 2025) reveals: labeled "
        "gateway volume is more concentrated in large transfers ($1M+: 84.3 percent of labeled volume "
        "versus 66.7 percent of unlabeled, a 17.6 pp gap), and labeled gateways exhibit six times the "
        "median counterparty degree (261 versus 41), consistent with the unlabeled population being "
        "structurally more P2P-like than gateway-like. 39 percent of the top 500 unlabeled addresses "
        "do exhibit gateway-like behavioral signatures (flow symmetry >= 0.9 and counterparty degree "
        ">= 100), confirming that some unlabeled volume is exchange-like, but this does not dominate "
        "the distribution."
    ))
    add_exhibit(doc, "exhibit_sb1_unlabeled_degree.png",
        "Exhibit SB-1. Behavioral signatures of unlabeled addresses. Panel A: counterparty degree "
        "versus total volume. Panel B: flow symmetry density. Labeled gateways cluster at higher "
        "degree (median 261 versus 41). Source: Dune Analytics, 6-month sample.")
    add_exhibit(doc, "exhibit_sb2_size_distribution.png",
        "Exhibit SB-2. Transfer size distribution, labeled versus unlabeled addresses. Panel A: "
        "share of volume by transfer size bucket. Panel B: share of transfer count. "
        "Source: Dune Analytics, Jul 2024-Jan 2025, Ethereum USDC + USDT.")

def build_s14_tokenized_fund(doc):
    add_heading(doc, "Tokenized Fund Product Assessment", level=1)
    add_para(doc, (
        "As tokenized Treasury fund products grow, a natural question is whether these products "
        "could enter the stablecoin reserve stack or serve as functional substitutes for traditional "
        "money market fund shares. We assess three leading products against nine categories of "
        "minimum viable economic parity (MVEP) with traditional money market instruments."
    ))
    headers = ["Category", "BUIDL (BlackRock)", "BENJI (Franklin T.)", "OUSG (Ondo)"]
    rows = [
        ["Redemption guarantee", "Pass", "Pass", "Ambiguous"],
        ["Regulatory licensing", "Ambiguous", "Pass", "Ambiguous"],
        ["Segregation of assets", "Pass", "Pass", "Pass"],
        ["Qualified custodian", "Pass", "Pass", "Pass"],
        ["Daily NAV transparency", "Pass", "Pass", "Pass"],
        ["AML/KYC gateway", "Pass", "Pass", "Pass"],
        ["Smart contract audit", "Ambiguous", "Ambiguous", "Pass"],
        ["Interoperability", "Pass", "Ambiguous", "Pass"],
        ["Regulatory clarity", "Pass", "Pass", "Ambiguous"],
        ["Score", "7/9", "7/9", "6/9"],
    ]
    add_simple_table(doc, headers, rows)
    add_source(doc, (
        "Table G1. Minimum Viable Economic Parity (MVEP) assessment. BlackRock BUIDL AUM ~$2.5B "
        "(Nov 2025); Franklin BENJI ~$844M; Ondo OUSG ~$770M. Source: Author's analysis of "
        "product documentation, SEC filings, and audit reports as of January 2026."
    ))
    add_para(doc, (
        "All three products pass on core institutional requirements: segregation of assets, "
        "qualified custodian, daily NAV transparency, and AML/KYC gateways. The residual ambiguity "
        "clusters in two areas. First, regulatory licensing: BUIDL operates as an unregistered "
        "private placement under Regulation D, while BENJI is the only product with full SEC "
        "registration as a mutual fund. Second, smart contract auditability: neither BUIDL nor "
        "BENJI has published comprehensive, publicly available independent smart contract audits, "
        "while OUSG has undergone multiple audits."
    ))
    add_para(doc, (
        "The gateway implications are significant. Tokenized fund products create a new class of "
        "potential Tier 1 gateway in which the fund administrator and transfer agent combines "
        "reserve management with token issuance, collapsing the issuer-gateway distinction that "
        "the CLII currently separates."
    ))

def build_s15_credit_migration(doc):
    add_heading(doc, "The Credit Migration Channel", level=1)
    add_para(doc, (
        "The deposit-displacement mechanism described in Section IV.E has a second-order consequence "
        "that connects stablecoin growth directly to the structure of credit intermediation. "
        "Stablecoin reserves are, by design, narrow: the GENIUS Act mandates one-to-one backing "
        "in cash, Treasury bills, and repo, none of which fund loans. Every dollar of deposits that "
        "migrates into stablecoin reserves is a dollar removed from the fractional-reserve lending channel."
    ))
    add_para(doc, (
        "The credit migration channel operates through four distinct mechanisms: "
        "(A) deposit displacement itself: the money object moves from bank liabilities to stablecoin "
        "reserves, reducing the deposit base available for lending; "
        "(B) credit supply relocation: the marginal loan shifts from bank to nonbank balance sheets; "
        "(C) funding fragility: replacement credit is funded by wholesale markets and securitization "
        "rather than insured deposits; and "
        "(D) backstop ambiguity: nonbank credit intermediaries may lack access to Fed lending facilities."
    ))
    add_para(doc, (
        "The closest precedent is the post-2008 restructuring of small-business credit documented "
        "by Gopal and Schnabl (2022), who show that regulatory constraints on banks drove a wholesale "
        "migration of lending activity to nonbank finance companies and fintech originators. Three "
        "features of their findings map onto the stablecoin context: the regulatory perimeter predicted "
        "behavior under constraint, replacement credit shifted toward shorter maturities and heavier "
        "collateral, and the speed of nonbank entry was striking."
    ))
    add_para(doc, (
        "Congress's yield prohibition on payment stablecoins responds in part to this deposit-displacement "
        "concern. The prohibition echoes Regulation Q's mid-twentieth-century interest-rate ceilings. "
        "A parallel development -- tokenized deposits -- offers an inside-the-perimeter alternative: "
        "bank deposits represented as tokens that remain on the bank's balance sheet, inherit deposit "
        "insurance, and preserve the fractional-reserve lending channel."
    ))
    add_para(doc, (
        "For monetary policy, the credit migration channel implies that the transmission externality "
        "operates on two margins simultaneously: the interest-rate channel is attenuated if displaced "
        "liquidity circulates in venues unconstrained by reserve requirements, and the credit channel "
        "is restructured if deposit displacement pushes credit supply toward more procyclical nonbank "
        "balance sheets. At current scale ($305 billion), these effects are negligible. At the $1-2 "
        "trillion range that current growth rates imply by the end of the decade, the deposit-displacement "
        "effect could rival the post-2008 bank lending contraction."
    ))

def build_s16_eurodollar(doc):
    add_heading(doc, "Appendix J: The Eurodollar Analogy and Its Limits", level=1)
    add_para(doc, (
        "The offshore USDT ecosystem shares three defining features with the eurodollar system that "
        "emerged in the 1960s: dollar instruments circulating outside U.S. regulatory infrastructure, "
        "different monetary transmission properties than their domestic counterparts, and a structural "
        "challenge to the Fed's ability to monitor and backstop dollar liquidity globally. As with "
        "eurodollars, the growth is demand-driven (emerging-market dollar demand) rather than "
        "policy-designed, and the regulatory response has lagged the market's development by years."
    ))
    add_para(doc, (
        "The analogy is imprecise in three important respects. First, eurodollars create money through "
        "fractional-reserve lending while stablecoins are fully reserved; the credit-creation multiplier "
        "that gave eurodollar markets their systemic significance is absent in stablecoin markets by "
        "design. Second, eurodollar banks retain access to lender-of-last-resort facilities through "
        "central bank swap lines while stablecoin issuers have no equivalent backstop; the Fed's "
        "swap-line architecture covers eurodollar stress but not stablecoin stress. Third, eurodollar "
        "deposit creation is driven by credit demand while stablecoin issuance is driven by fiat "
        "inflow from token purchasers, a distinction that affects the elasticity of supply to monetary "
        "conditions."
    ))
    add_para(doc, (
        "Despite these differences, the regulatory challenge is structurally parallel: monitoring dollar "
        "usage beyond the supervisory perimeter. The eurodollar experience suggests that once dollar "
        "instruments achieve sufficient offshore scale, they become self-reinforcing through network "
        "effects and difficult to repatriate within the regulatory perimeter without disrupting the "
        "economic functions they serve. The stablecoin routing infrastructure documented in this paper "
        "represents an early-stage version of this dynamic, with the additional complication that the "
        "\"offshore\" dimension is defined by blockchain and gateway choice rather than geographic "
        "jurisdiction."
    ))

# ── Main builder ─────────────────────────────────────────────

def main():
    print("=" * 60)
    print("BUILDING ONLINE SUPPLEMENT v42")
    print("=" * 60)

    doc = Document()
    set_style(doc)

    # Title page
    add_para(doc, "Online Supplement", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "", size=6)
    add_para(doc, (
        "Routing the Dollar: Gateway Infrastructure, Monetary Policy Transmission, "
        "and the Dollar's International Functions in Digital Markets"
    ), bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "", size=6)
    add_para(doc, "[Author Name]", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "[Institutional Affiliation]", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_para(doc, "", size=6)
    add_para(doc, "March 2026", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "", size=6)
    add_para(doc, (
        "Prepared for the Fifth Conference on International Roles of the U.S. Dollar, "
        "Board of Governors of the Federal Reserve System and Federal Reserve Bank of New York, "
        "June 22-23, 2026."
    ), size=10, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

    doc.add_page_break()

    # Build all sections
    print("  S1: Introduction...")
    build_s1_intro(doc)

    print("  S2: Exhibit Inventory (Appendix A)...")
    doc.add_page_break()
    build_s2_exhibit_inventory(doc)

    print("  S3: Tables F1-F3...")
    doc.add_page_break()
    build_s3_tables_f1_f3(doc)

    print("  S4: Extended FOMC Event Study...")
    doc.add_page_break()
    build_s4_fomc_extended(doc)

    print("  S5: Unit Root Tests (ADF/KPSS)...")
    build_s5_unit_root(doc)

    print("  S6: Weak Exogeneity Test...")
    doc.add_page_break()
    build_s6_weak_exogeneity(doc)

    print("  S7: FEVD and Persistence Profile...")
    build_s7_fevd_persistence(doc)

    print("  S8: SVB All 12 Entities...")
    doc.add_page_break()
    build_s8_svb_all_entities(doc)

    print("  S9: Placebo Analysis...")
    doc.add_page_break()
    build_s9_placebo(doc)

    print("  S10: Day-of-Week Structure...")
    doc.add_page_break()
    build_s10_day_of_week(doc)

    print("  S11: No-Freeze CLII Robustness...")
    doc.add_page_break()
    build_s11_no_freeze(doc)

    print("  S12: Unlabeled Address Characterization...")
    doc.add_page_break()
    build_s12_unlabeled(doc)

    print("  S14: Tokenized Fund Assessment...")
    doc.add_page_break()
    build_s14_tokenized_fund(doc)

    print("  S15: Credit Migration Channel...")
    doc.add_page_break()
    build_s15_credit_migration(doc)

    print("  S16: Eurodollar Analogy (Appendix J)...")
    doc.add_page_break()
    build_s16_eurodollar(doc)

    # ── Language sync ────────────────────────────────────────
    print("\n  Language sync...")
    # We need to iterate through all paragraphs and fix text
    total_fixes = 0
    for para in doc.paragraphs:
        for run in para.runs:
            original = run.text
            fixed, n = language_sync(original)
            fixed = fix_style(fixed)
            if fixed != original:
                run.text = fixed
                total_fixes += n

    # Also check table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        original = run.text
                        fixed, n = language_sync(original)
                        fixed = fix_style(fixed)
                        if fixed != original:
                            run.text = fixed
                            total_fixes += n

    print(f"  Total language fixes: {total_fixes}")

    # ── Style check ──────────────────────────────────────────
    print("\n  Style check...")
    all_text = ""
    for para in doc.paragraphs:
        all_text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text + "\n"

    checks = style_check(all_text)
    print(f"    Em-dashes: {checks['em_dashes']}")
    print(f"    Author's (singular): {checks['author_singular']}")
    print(f"    Authors' (plural): {checks['authors_plural']}")
    print(f"    'stress event': {checks['stress_event']}")

    # Check for Table B2 duplication
    has_table_b2 = 'Table B2' in all_text
    print(f"    Table B2 in supplement: {'YES (needs removal)' if has_table_b2 else 'No (correct)'}")

    # ── Save ─────────────────────────────────────────────────
    doc.save(str(OUTPUT))
    word_count = len(all_text.split())
    print(f"\n  Saved: {OUTPUT}")
    print(f"  Word count: ~{word_count}")

    # ── Resolution report ────────────────────────────────────
    print("\n" + "=" * 60)
    print("SUPPLEMENT SYNC REPORT")
    print("=" * 60)

    resolutions = [
        ("1", "Front matter definition", "Introduction" in all_text and "companion" in all_text.lower()),
        ("2", "Omitted appendix letters (A, J)", "Appendix A" in all_text and "Appendix J" in all_text),
        ("3", "No-freeze CLII robustness", "no-freeze" in all_text.lower() or "No-Freeze" in all_text),
        ("4", "ADF and KPSS testing", "ADF" in all_text and "KPSS" in all_text),
        ("5", "Weak exogeneity test", "weak exogeneity" in all_text.lower() or "Weak Exogeneity" in all_text),
        ("6", "FEVD and persistence profile", "FEVD" in all_text and "persistence" in all_text.lower()),
        ("7", "All 12 gateway entities", "12 Entities" in all_text or "All 12" in all_text),
        ("8", "Placebo windows", "placebo" in all_text.lower() and "50" in all_text),
        ("9", "Full robustness summary", "Robustness" in all_text and "F2" in all_text),
        ("10", "Day-of-week, no-freeze, unlabeled", all(x in all_text.lower() for x in ["day-of-week", "unlabeled"])),
        ("11", "Tables F1-F3 and FOMC", all(x in all_text for x in ["Table F1", "Table F2", "Table F3"])),
    ]

    all_pass = True
    for num, desc, resolved in resolutions:
        status = "RESOLVES" if resolved else "MISSING"
        if not resolved:
            all_pass = False
        print(f"  Reference {num}: {status} -- {desc}")

    print(f"\n  Language fixes: {total_fixes} stale strings replaced")
    print(f"  Style: em-dashes {checks['em_dashes']}, "
          f"author attribution {'singular' if checks['authors_plural'] == 0 else 'MIXED'}")
    print(f"\n  OVERALL: {'PASS' if all_pass else 'ISSUES FOUND'}")
    print("=" * 60)

if __name__ == "__main__":
    main()
